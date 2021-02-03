import os
import shutil
import docx

from datetime import datetime
from openpyxl import load_workbook
from parse import TBC

RESOURCES_FOLDER = os.path.abspath('resources')


def get_new_filename(record, document_name, file_extension):
    """Generates the filename of an output document

    Args:
        record(Record): A Record object
        document_name(str): The name of the document that is being exported
        file_extension(str): The file extension of the documents to copy

    Returns:
        str: The new filename

    """
    now = datetime.now()
    curr_year = now.year
    curr_month = now.strftime('%b')

    new_filename = (
        f'{record.client.last_name}, '
        f'{record.client.first_name} - '
        f'{document_name} - '
        f'{curr_year} '
        f'{curr_month}'
        f'.{file_extension}'
    )

    return new_filename


def copy_resources_to_export(record, export_folder, file_extension):
    """Copies all files with the given extension from the resources folder to the export folder

    Args:
        record(Record): A Record object
        export_folder(str): The absolute path of the folder to export to
        file_extension(str): The file extension of the documents to copy

    Returns:
        list(str): The list of destination file absolute paths

    """
    paths = []
    for item in os.listdir(RESOURCES_FOLDER):
        if not item.endswith(f'.{file_extension}'):
            continue

        document_name = item[:item.index(f'.{file_extension}')]
        src = os.path.join(RESOURCES_FOLDER, item)
        dst = os.path.join(export_folder, get_new_filename(record, document_name, file_extension))
        shutil.copyfile(src, dst)

        paths.append(dst)

    return paths


def word_export(record, export_folder):
    """Exports the data in a Record object into all of the output word documents

    Args:
        record(Record): A Record object
        export_folder(str): The absolute path of the folder to export to

    Returns:
        None

    """
    placeholder_to_val = {
        '[full_name]': record.client.full_name,
        '[dob]': record.client.dob,
        '[gender]': record.client.gender,
        '[address]': str(record.client.address),
        '[house_number]': record.client.address.house_number,
        '[street]': record.client.address.street,
        '[suburb]': record.client.address.suburb,
        '[state]': record.client.address.state,
        '[home_phone_number]': record.client.home_phone_number,
        '[mobile_phone_number]': record.client.mobile_phone_number,
        '[email_address]': record.client.email_address,
        '[ndis_number]': record.client.ndis_number,
        '[plan_start_date]': record.plan.start_date,
        '[plan_end_date]': record.plan.end_date,
        '[core_supports_categories]': record.supports['Core'].categories,
        '[capacity_building_supports_categories]': record.supports['Capacity Building'].categories,
        '[capital_supports_categories]': record.supports['Capital'].categories,
        '[core_supports_total]': record.supports['Core'].total,
        '[capacity_building_supports_total]': record.supports['Capacity Building'].total,
        '[capital_supports_total]': record.supports['Capital'].total,
        '[support_coordination_hours]': record.support_coordination_hours
    }

    goals = []
    for value in record.supports.values():
        if value.goals != TBC:
            for goal in value.goals:
                goals.append(goal)
        else:
            goals.append('')

    goals.extend(['' for _ in range(12 - len(goals))])

    def search_and_replace(paragraph):
        for placeholder in placeholder_to_val.keys():
            if placeholder in paragraph.text:
                value = placeholder_to_val[placeholder]
                if 'categories' in placeholder:
                    if value == TBC:
                        string = TBC
                    else:
                        string = ''
                        for category in value:
                            string += f'{category[0]}: {category[1]}\n'
                else:
                    string = value

                paragraph.text = paragraph.text.replace(
                    placeholder,
                    string
                )
            elif '[goal]' in paragraph.text:
                paragraph.text = paragraph.text.replace('[goal]', goals[0])
                goals.pop(0)
            elif '[sc1]' in paragraph.text:
                if record.support_coordination_management_type.lower() == 'ndia-managed':
                    paragraph.text = '   X'
                else:
                    paragraph.text = ''
            elif '[sc2]' in paragraph.text:
                if record.support_coordination_management_type.lower() == 'self-managed':
                    paragraph.text = '   X'
                else:
                    paragraph.text = ''

    paths = copy_resources_to_export(record, export_folder, 'docx')
    for path in paths:
        doc = docx.Document(path)
        for paragraph in doc.paragraphs:
            search_and_replace(paragraph)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        search_and_replace(paragraph)

        doc.save(path)


def excel_export(record, export_folder='', optional_xml_path=''):
    """Exports the data in a Record object into all of the output excel documents

    Args:
        record(Record): A Record object
        export_folder(str): The absolute path of the folder to export to (optional)
        optional_xml_path(str): The path of an xml document to append data to if a new one
            should not be created (optional)

    Returns:
        None

    """
    data = (
        record.client.title,
        'CLIENT',
        record.client.first_name,
        record.client.last_name,
        record.client.home_phone_number,
        record.client.mobile_phone_number,
        record.client.gender,
        record.client.dob,
        record.client.email_address,
        record.additional_email_address,
        record.service_region_id,
        f'{record.client.address.house_number} {record.client.address.street}',
        '',
        record.client.address.suburb,
        record.client.address.state,
        record.client.address.postcode,
        record.client.ndis_number,
        record.plan.start_date,
        record.plan.end_date,
        'Support Coordination Client',
        'Inactive',
        TBC,
        TBC,
        TBC,
        TBC,
        'off',
        TBC,
        TBC,
        TBC,
        TBC,
        'Yes',
        TBC,
        TBC,
        'No',
        'Yes',
        'No',
        TBC,
        TBC,
        'None',
        TBC,
        TBC,
        'Megan King',
        'Megan King'
    )

    if export_folder:
        path = copy_resources_to_export(record, export_folder, 'xlsx')[0]
    else:
        path = optional_xml_path

    wb = load_workbook(filename=path)
    ws = wb.active
    ws.append(data)
    wb.save(path)
